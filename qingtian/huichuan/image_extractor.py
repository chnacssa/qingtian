"""汇川图片提取器 — PDF/DOCX 嵌入图片 → Layer 1 存储

Phase 2 模块。从 PDF 和 DOCX 文件中提取嵌入图片，保存到 Layer 1 存储。

与 _extract_pdf()/_extract_docx() (ingest.py) 的关系：
  - _extract_pdf/_extract_docx 只提取文本（已有功能，不改动）
  - 本模块提取图片（新增功能）
  两者在 ingest_file() 中串行调用：先提取文本，再提取图片。
"""

from __future__ import annotations

import asyncio
import hashlib
import logging
import os
from datetime import date

logger = logging.getLogger("huichuan.image_extractor")

# ── 边界常量 ────────────────────────────────────────────

MAX_IMAGES = 50                # 单文件最大提取图片数
MAX_IMAGE_SIZE = 10 * 1024 * 1024  # 单张最大 10MB


class ImageRecord:
    """单张提取图片的记录。

    Attributes:
        data: 图片原始字节
        fmt: 格式名 (png/jpg/jp2/bmp/…)
        page_num: PDF 页码（仅 PDF）
        image_index: 文件内图片序号
        source_sheet: Excel Sheet 名（仅 XLSX）
        size: 字节数
        sha256: SHA256 指纹
        width/height: 像素尺寸
        storage_path: Layer 1 存储路径
        context_before/after: 前后文各 200 字
    """
    def __init__(self, data: bytes, fmt: str, *, page_num: int = 0,
                 image_index: int = 0, source_sheet: str = "",
                 context_before: str = "", context_after: str = ""):
        self.data = data
        self.fmt = fmt
        self.page_num = page_num
        self.image_index = image_index
        self.source_sheet = source_sheet
        self.context_before = context_before[:200]
        self.context_after = context_after[:200]
        self.size = len(data)
        self.sha256 = hashlib.sha256(data).hexdigest()
        self.storage_path = ""
        # 同步解析尺寸（图片通常 < 10MB，不阻塞事件循环）
        self.width = 0
        self.height = 0
        self._resolve_dims()

    def _resolve_dims(self) -> None:
        try:
            from PIL import Image
            import io
            img = Image.open(io.BytesIO(self.data))
            self.width, self.height = img.size
        except Exception:
            pass


async def extract_from_pdf(pdf_data: bytes, storage_base: str,
                           file_id: str) -> list[ImageRecord]:
    """从 PDF 提取嵌入图片。

    内部用 asyncio.to_thread 避免阻塞事件循环。
    pdfplumber page.images 返回 dict，含 stream/filter/name 等元信息。

    Args:
        pdf_data: PDF 文件原始字节
        storage_base: Layer 1 根目录
        file_id: 文件注册表 ID（用于子目录命名）

    Returns:
        list[ImageRecord] — 提取到的图片列表（可能为空）
    """
    import pdfplumber
    import io

    def _extract() -> list[ImageRecord]:
        records: list[ImageRecord] = []
        page_texts: list[str] = []

        with pdfplumber.open(io.BytesIO(pdf_data)) as pdf:
            page_texts = [(p.extract_text() or "") for p in pdf.pages]

            for pi, page in enumerate(pdf.pages):
                if len(records) >= MAX_IMAGES:
                    break
                for img in page.images:
                    if len(records) >= MAX_IMAGES:
                        break
                    img_data = img.get("stream", img.get("data"))
                    if not img_data:
                        continue
                    raw = img_data.get_data() if hasattr(img_data, "get_data") else img_data
                    if not isinstance(raw, bytes) or len(raw) > MAX_IMAGE_SIZE:
                        continue
                    fmt = _detect_image_format(img_data, img)
                    before = page_texts[pi][:200] if pi < len(page_texts) else ""
                    after = page_texts[pi + 1][:200] if pi + 1 < len(page_texts) else ""
                    record = ImageRecord(
                        raw, fmt,
                        page_num=pi,
                        image_index=len(records),
                        context_before=before,
                        context_after=after,
                    )
                    record.storage_path = _save_image(raw, fmt, storage_base, file_id, len(records))
                    records.append(record)

        return records

    return await asyncio.to_thread(_extract)


async def extract_from_docx(docx_data: bytes, storage_base: str,
                            file_id: str) -> list[ImageRecord]:
    """从 DOCX 提取嵌入图片。

    通过 python-docx 的 related_parts 访问嵌入的图片 blob。
    每个图片关联到所在段落，提取前后段落作为上下文。

    Args:
        docx_data: DOCX 文件原始字节
        storage_base: Layer 1 根目录
        file_id: 文件注册表 ID

    Returns:
        list[ImageRecord] — 提取到的图片列表（可能为空）
    """
    from docx import Document
    import io

    def _extract() -> list[ImageRecord]:
        records: list[ImageRecord] = []
        doc = Document(io.BytesIO(docx_data))
        paragraphs = [p.text for p in doc.paragraphs]

        # 收集所有 inline 图片
        for pi, para in enumerate(doc.paragraphs):
            if len(records) >= MAX_IMAGES:
                break
            for run in para.runs:
                if len(records) >= MAX_IMAGES:
                    break
                # 检查 run 是否包含 drawing（图片标记）
                draw_elems = run._element.findall(
                    ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
                )
                if not draw_elems:
                    continue
                for blip in run._element.iter(
                    "{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
                ):
                    if len(records) >= MAX_IMAGES:
                        break
                    embed_id = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if not embed_id:
                        continue
                    img_part = doc.part.related_parts.get(embed_id)
                    if not img_part or not hasattr(img_part, "blob"):
                        continue
                    raw = img_part.blob
                    if not isinstance(raw, bytes) or len(raw) > MAX_IMAGE_SIZE:
                        continue
                    before = paragraphs[pi - 1][:200] if pi > 0 else ""
                    after = paragraphs[pi + 1][:200] if pi + 1 < len(paragraphs) else ""
                    fmt = detect_format_from_bytes(raw)
                    record = ImageRecord(
                        raw, fmt,
                        image_index=len(records),
                        context_before=before,
                        context_after=after,
                    )
                    record.storage_path = _save_image(raw, fmt, storage_base, file_id, len(records))
                    records.append(record)

        return records

    return await asyncio.to_thread(_extract)


def _save_image(data: bytes, fmt: str, storage_base: str,
                file_id: str, index: int) -> str:
    """保存图片到 Layer 1 (storage/{yyyy}/{mm}/images/{file_id}/)。

    此函数被 extract_from_pdf/extract_from_docx 在线程内调用，
    因此可直接做同步 I/O，无需额外 to_thread。

    Returns:
        图片文件的完整路径
    """
    today = date.today()
    img_dir = os.path.join(storage_base, str(today.year), f"{today.month:02d}",
                          "images", file_id)
    os.makedirs(img_dir, exist_ok=True)
    ext = f".{fmt}" if fmt else ".png"
    path = os.path.join(img_dir, f"{index}{ext}")
    with open(path, "wb") as f:
        f.write(data)
    return path


def _detect_image_format(stream_obj, img_meta: dict) -> str:
    """从 pdfplumber stream/filter/name 推断图片格式。

    优先级: filter 名 > 魔数 > name 扩展名 > 兜底 png。
    """
    # 1. filter 推断
    cf = img_meta.get("filter") or ""
    cf_str = cf if isinstance(cf, str) else (cf[0] if isinstance(cf, (list, tuple)) and cf else "")
    cf_lower = cf_str.lower()
    if "jpeg" in cf_lower or "dct" in cf_lower:
        return "jpg"
    if "jp2" in cf_lower:
        return "jp2"
    if "ccitt" in cf_lower:
        return "tiff"

    # 2. 魔数推断
    if hasattr(stream_obj, "get_data"):
        try:
            head = stream_obj.get_data()[:12]
            if isinstance(head, bytes):
                if head[:2] == b"\xff\xd8":
                    return "jpg"
                if head[:8] == b"\x89PNG\r\n\x1a\n":
                    return "png"
                if head[:4] == b"II\x2a\x00" or head[:4] == b"MM\x00\x2a":
                    return "tiff"
        except Exception:
            pass

    # 3. name 扩展名兜底
    name = img_meta.get("name", "")
    if "." in name:
        ext = name.rsplit(".", 1)[-1].lower()
        if ext in ("png", "jpg", "jpeg", "jp2", "bmp", "tiff", "gif"):
            return "jpg" if ext == "jpeg" else ext
        if ext == "jpeg":
            return "jpg"

    return "png"


def detect_format_from_bytes(data: bytes) -> str:
    """从原始字节魔数推断图片格式（用于 DOCX blob 等无元数据的场景）。"""
    if len(data) < 4:
        return "png"
    if data[:2] == b"\xff\xd8":
        return "jpg"
    if data[:8] == b"\x89PNG\r\n\x1a\n":
        return "png"
    if data[:4] == b"GIF8":
        return "gif"
    if data[:4] == b"RIFF":
        return "webp"
    if data[:4] in (b"II\x2a\x00", b"MM\x00\x2a"):
        return "tiff"
    return "png"
