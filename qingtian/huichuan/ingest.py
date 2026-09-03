"""汇川 LLM 摄入管道 — 文件→文本→LLM编译→入库

Phase 3 核心模块。LLM 在摄入时编译知识，而非查询时拼接。

边界约束:
  - text 空 → {"entries":0, "error":"empty"}
  - text > MAX_CHUNK_CHARS → 截断前 MAX_CHUNK_CHARS 字符 + 记录警告
  - LLM 调用失败 → {"entries":0, "error":str(e)}
  - 单文档最多生成 MAX_ENTRIES_PER_DOC 条
  - 每条 entry 写入前做 sanitize(PII)
"""

from __future__ import annotations

import asyncio
import hashlib
import io
import json
import logging
import os
import uuid
from datetime import date, datetime, timezone

from common.llm import llm_call_json
from huichuan.file_classifier import classify
from huichuan.image_extractor import extract_from_docx, extract_from_pdf
from huichuan.import_export import validate_entry
from huichuan.metadata_extractor import EXTRACTORS
from huichuan.sanitizer import sanitize

from . import config as kcfg

logger = logging.getLogger("huichuan.ingest")

# ── 边界常量 ────────────────────────────────────────────

MAX_CHUNK_CHARS = 50000        # 单次 LLM 编译最大字符
MAX_ENTRIES_PER_DOC = 15       # 单文档最多生成 15 条知识
MIN_CONFIDENCE = 0.6            # LLM confidence 门槛（低于此值不入库）
COOLDOWN_SECONDS = 5           # 防并发 LLM 调用（留给调用方控制）

# entry_type 白名单 + 模糊匹配兜底（LLM 可能返回 fact/manual/report 等非标准值）
_ALLOWED_ENTRY_TYPES = frozenset({"entity", "concept", "comparison", "query", "source"})
# 9-1 可见性白名单：非白名单值统一回落 'enterprise'（原硬编码值，语义不变）
_ALLOWED_VISIBILITY = {"enterprise": "enterprise", "private": "private"}
_ENTRY_TYPE_FALLBACK = {
    "fact": "entity", "manual": "source", "spec": "concept",
    "policy": "concept", "report": "source", "meeting": "query",
    "other": "entity", "analysis": "comparison", "summary": "concept",
}


def _normalize_entry_type(raw: str) -> str:
    """将 LLM 返回的 entry_type 规范化为 DB 约束白名单内的值。"""
    t = (raw or "").strip().lower()
    if t in _ALLOWED_ENTRY_TYPES:
        return t
    if t in _ENTRY_TYPE_FALLBACK:
        return _ENTRY_TYPE_FALLBACK[t]
    return "entity"

# ── LLM Prompt ──────────────────────────────────────────

INGEST_PROMPT = """分析以下文档，提取结构化知识。返回 JSON:
{
  "entries": [
    {"title":"...","domain":"...","entry_type":"entity|concept","content":"摘要",
     "tags":["..."],"confidence":0.0-1.0,
     "related_to":["已有知识标题"],"contradicts":[{"title":"...","reason":"..."}]}
  ],
  "summary":"一句话总结"
}
规则:
- 每个独立事实/概念/数据一条 entry
- domain 优先使用文档来源分类（如"采购部""总裁办""财务部"→对应domain），
  而非内容中提到的行业分类。来源不明时再根据内容判断领域
- confidence: 原文明确=1.0, 推断=0.7, 不确定=0.5
- 发现与已有知识的矛盾时标出 contradicts
"""


def _now_str() -> str:
    """返回 ISO 字符串（仅用于 JSONB metadata，非 asyncpg TIMESTAMPTZ 参数）。"""
    return datetime.now(timezone.utc).isoformat()


async def ingest_text(
    conn,
    text: str,
    source: str = "manual",
    original_filename: str = "",
    storage_path: str = "",
    schema: str = "huichuan",
    visibility: str = "enterprise",
) -> dict:
    """摄入文本 → LLM 编译 → 批量入库。

    visibility（9-1 参数化，此前硬编码 'enterprise'）：白名单外的值一律回落
    'enterprise'（保守面：不产生新公开面，仅允许显式收紧到 private）。

    边界约束:
      - text 空 → 返回 {"entries":0, "error":"empty"}
      - text > MAX_CHUNK_CHARS → 截断前 MAX_CHUNK_CHARS 字符 + 记录警告
      - LLM 调用失败 → 返回 {"entries":0, "error":str(e)}
      - 单文档最多生成 MAX_ENTRIES_PER_DOC 条
      - 每条 entry 写入前做 sanitize(PII)

    Args:
        conn: asyncpg connection
        text: 待编译的原始文本
        source: 来源标识 ("manual" | "connector:{name}" | "feishu" | "api")
        original_filename: 原始文件名（溯源）
        storage_path: Layer 1 存储路径
        schema: 数据库 schema 名

    Returns:
        {"entries": N, "summary": "...", "ingested_at": "...", "knowledge_ids": [...]}
    """
    if not text or not text.strip():
        return {"entries": 0, "error": "empty text", "ingested_at": _now_str()}

    text = text.strip()[:MAX_CHUNK_CHARS]
    if len(text) >= MAX_CHUNK_CHARS:
        logger.warning("Text truncated to %d chars for ingest", MAX_CHUNK_CHARS)

    # DeepSeek 缓存优化：固定 prompt 放 system 角色以命中缓存
    user_prompt = f"文档来源: {source}\n文档内容:\n{text}"

    result = await llm_call_json(
        user_prompt,
        caller="huichuan.ingest",
        default={"entries": [], "summary": ""},
        timeout=120,
        max_tokens=3000,
        system_prompt=INGEST_PROMPT,
    )

    if not result or not result.get("entries"):
        return {"entries": 0, "error": "LLM returned empty", "ingested_at": _now_str()}

    entries = result["entries"][:MAX_ENTRIES_PER_DOC]
    ingested = []

    # 事务包裹：LLM 返回的 N 条 entry 全部成功或全部回滚
    if conn is None:
        logger.error("ingest_text: conn is None, cannot proceed with transaction")
        return {"entries": 0, "error": "no database connection", "ingested_at": _now_str()}
    async with conn.transaction():
        skipped_low_conf = 0
        skipped_validation = 0

        for entry in entries:
            # ── 质量门 1: LLM confidence 门槛 ────────────────
            confidence = entry.get("confidence", 0.7)
            if confidence < MIN_CONFIDENCE:
                skipped_low_conf += 1
                logger.debug("Entry '%s' skipped: confidence %.2f < %.2f",
                             entry.get("title", "?")[:40], confidence, MIN_CONFIDENCE)
                continue

            # Sanitize PII（复用 sanitizer.py）
            content = sanitize(entry.get("content", ""), level="erp_to_ingest")
            title = sanitize(entry.get("title", ""), level="erp_to_ingest")[:256]

            # ── 质量门 2: 准入校验（复验 PII/长度/黑名单）────
            violation = validate_entry(title, content, entry.get("domain", "general"))
            if violation:
                skipped_validation += 1
                logger.warning("Entry '%s' failed validation: %s", title[:60], violation)
                continue

            row = await conn.fetchrow(
                f"""INSERT INTO {schema}.knowledge_entries
                   (title, domain, content, tags, visibility, entry_type,
                    original_filename, original_storage_path, metadata, quality)
                   VALUES ($1,$2,$3,$4,$10,$5,$6,$7,
                           $8::jsonb, $9)
                   RETURNING knowledge_id""",
                title,
                entry.get("domain", "general"),
                content,
                entry.get("tags", []),
                _normalize_entry_type(entry.get("entry_type", "entity")),
                original_filename,
                storage_path,
                json.dumps(
                    {
                        "confidence": confidence,
                        "source": source,
                        "ingested_at": _now_str(),
                        "related_to": entry.get("related_to", []),
                        "contradictions": entry.get("contradicts", []),
                    },
                    ensure_ascii=False,
                ),
                max(1, min(int(confidence * 5), 5)),
                _ALLOWED_VISIBILITY.get(visibility, "enterprise"),
            )
            if row:
                ingested.append(str(row["knowledge_id"]))

    # ── 文件注册表同步 ────────────────────────────────
    if storage_path and original_filename:
        await _upsert_file_registry(
            conn, storage_path, original_filename,
            entries_total=len(ingested),
            schema=schema,
        )

    logger.info(
        "Ingested %d entries from '%s' (skipped: %d low-confidence, %d validation)",
        len(ingested),
        original_filename or "text",
        skipped_low_conf,
        skipped_validation,
    )
    return {
        "entries": len(ingested),
        "summary": result.get("summary", ""),
        "ingested_at": _now_str(),
        "knowledge_ids": ingested,
        "skipped_low_confidence": skipped_low_conf,
        "skipped_validation": skipped_validation,
    }


async def ingest_file(
    conn,
    file_bytes: bytes,
    filename: str,
    source: str = "api",
    storage_base: str = "/opt/qingtian/huichuan/storage",
    schema: str = "huichuan",
) -> dict:
    """摄入文件 → 解析文本 → ingest_text 入库。

    支持格式: .txt, .md, .json, .csv, .pdf, .docx, .xlsx

    Args:
        conn: asyncpg connection
        file_bytes: 文件原始字节
        filename: 原始文件名（含扩展名）
        source: 来源标识
        storage_base: Layer 1 存储根目录
        schema: 数据库 schema 名

    Returns:
        ingest_text 的返回值 + storage_path
    """
    ext = os.path.splitext(filename)[1].lower()
    today = date.today()

    # Layer 1: 保存原始文件（I/O 走线程池避免阻塞事件循环）
    storage_dir = os.path.join(storage_base, str(today.year), f"{today.month:02d}")
    await asyncio.to_thread(os.makedirs, storage_dir, exist_ok=True)

    file_id = uuid.uuid4().hex
    storage_path = os.path.join(storage_dir, f"{file_id}{ext}")

    # 文件 I/O — 在线程池执行（避免阻塞事件循环）
    await asyncio.to_thread(_write_file, storage_path, file_bytes)

    # 提取文本
    text = await _extract_text(file_bytes, filename)
    file_sha256 = hashlib.sha256(file_bytes).hexdigest()
    file_size = len(file_bytes)

    # ── 格式分类 + 分支调度 ──────────────────────────────
    fc = classify(file_bytes, filename)

    if not fc.processable:
        # 不可处理格式 → 只存元数据，不塞 LLM
        meta: dict = {"mime": fc.mime, "category": fc.category}
        if fc.category in EXTRACTORS:
            try:
                extracted = await asyncio.to_thread(
                    EXTRACTORS[fc.category], file_bytes, filename,
                )
                meta.update(extracted)
            except Exception as ex:
                logger.debug("Metadata extraction failed: %s", ex)
        meta["unknown_format"] = fc.fmt == "unknown"

        await _upsert_file_registry(
            conn, storage_path, filename,
            sha256=file_sha256, file_size=file_size,
            status="metadata_only",
            schema=schema,
        )
        await conn.execute(
            f"UPDATE {schema}.file_registry SET metadata = $1::jsonb "
            f"WHERE storage_path = $2",
            json.dumps(meta, ensure_ascii=False),
            storage_path,
        )
        return {
            "entries": 0, "error": None,
            "warning": f"Unprocessable format: {fc.mime}",
            "storage_path": storage_path,
            "status": "metadata_only",
            "future_processable": fc.future_processable,
            "ingested_at": _now_str(),
        }

    if ext == ".xlsx" and kcfg.get_excel_sheet_independent():
        # XLSX Sheet 独立编译分支（Phase 3）：每 Sheet 单独入库
        from huichuan.excel_processor import xlsx_to_entries, process_xlsx

        # P2 (R11): 只全量解析一次（含图片），把解析结果复用于 LLM 编译 + 图片注册，
        # 不再对同一文件二次全量解析（此前 xlsx_to_entries 内部 + 此处各解析一遍）。
        sheets = await process_xlsx(file_bytes, storage_base, file_id)
        xlsx_results = await xlsx_to_entries(
            file_bytes, source, filename, storage_path, conn, schema,
            storage_base=storage_base,
            sheets=sheets,
        )
        images_registered = 0
        for sheet in sheets:
            for img in sheet.images:
                await conn.execute(
                    f"""INSERT INTO {schema}.file_images
                        (file_id, source_type, source_sheet, image_index,
                         image_format, image_size, image_sha256, storage_path,
                         width, height)
                        VALUES ($1,'xlsx',$2,$3,$4,$5,$6,$7,$8,$9)""",
                    file_id, sheet.sheet_name, img.image_index, img.fmt,
                    img.size, img.sha256, img.storage_path,
                    img.width, img.height,
                )
                images_registered += 1
        await _upsert_file_registry(
            conn, storage_path, filename,
            sha256=file_sha256, file_size=file_size,
            entries_total=sum(r.get("entries", 0) for r in xlsx_results),
            schema=schema,
        )
        all_kids: list[str] = []
        all_summaries: list[str] = []
        for r in xlsx_results:
            kids = r.get("knowledge_ids") or []
            all_kids.extend(kids)
            summary = r.get("summary", "")
            if summary:
                all_summaries.append(summary)
        return {
            "entries": sum(r.get("entries", 0) for r in xlsx_results),
            "knowledge_ids": all_kids,
            "summary": " | ".join(all_summaries) if all_summaries else "",
            "xlsx_sheets": len(sheets),
            "images_registered": images_registered,
            "storage_path": storage_path,
            "ingested_at": _now_str(),
        }

    if not text:
        # 损坏文件注册
        await _upsert_file_registry(
            conn, storage_path, filename,
            sha256=file_sha256, file_size=file_size,
            status="corrupted",
            schema=schema,
        )
        return {
            "entries": 0,
            "error": f"Unable to extract text from {filename}",
            "storage_path": storage_path,
            "ingested_at": _now_str(),
        }

    result = await ingest_text(
        conn,
        text,
        source=source,
        original_filename=filename,
        storage_path=storage_path,
        schema=schema,
    )

    # ── 图片提取分支（PDF/DOCX 嵌入图片）───────────────
    images_registered = 0
    if ext in (".pdf", ".docx") and kcfg.get_image_extraction_enabled():
        try:
            max_images = kcfg.get_max_images_per_file()
            extractor_fn = extract_from_pdf if ext == ".pdf" else extract_from_docx
            images = await extractor_fn(file_bytes, storage_base, file_id)
            for img in images[:max_images]:
                await conn.execute(
                    f"""INSERT INTO {schema}.file_images
                        (file_id, source_type, page_num, image_index, image_format,
                         image_size, image_sha256, storage_path, width, height,
                         context_before, context_after)
                        VALUES ($1,$2,$3,$4,$5,$6,$7,$8,$9,$10,$11,$12)""",
                    file_id, ext[1:], img.page_num, img.image_index, img.fmt,
                    img.size, img.sha256, img.storage_path, img.width, img.height,
                    img.context_before, img.context_after,
                )
                images_registered += 1
        except Exception as e:
            logger.warning("Image extraction failed for %s: %s", filename, e)
    result["images_registered"] = images_registered

    # 补充 sha256/file_size（ingest_text 内部已调 _upsert_file_registry，
    # 此处做 update 补全文件指纹信息）
    await _upsert_file_registry(
        conn, storage_path, filename,
        sha256=file_sha256, file_size=file_size,
        entries_total=result.get("entries", 0),
        schema=schema,
    )
    return result


async def _upsert_file_registry(
    conn,
    storage_path: str,
    original_filename: str,
    sha256: str = "",
    file_size: int = 0,
    entries_total: int = 0,
    status: str = "active",
    schema: str = "huichuan",
) -> None:
    """文件注册表 upsert — 追踪 Layer 1 文件生命周期。"""
    await conn.execute(
        f"""INSERT INTO {schema}.file_registry AS fr
            (storage_path, original_filename, file_sha256, file_size,
             status, entries_total, updated_at)
            VALUES ($1, $2, $3, $4, $5, $6, NOW())
            ON CONFLICT (storage_path) DO UPDATE SET
              file_sha256 = COALESCE(NULLIF($3, ''), fr.file_sha256),
              file_size   = COALESCE(NULLIF($4, 0), fr.file_size),
              entries_total = COALESCE(NULLIF($6, 0), fr.entries_total),
              status      = CASE
                WHEN fr.status = 'revoked' THEN fr.status  -- 不覆盖 revoked
                ELSE $5
              END,
              updated_at  = NOW()""",
        storage_path, original_filename, sha256, file_size,
        status, entries_total,
    )


def _write_file(path: str, data: bytes) -> None:
    """同步写文件（供 asyncio.to_thread 调用）。"""
    with open(path, "wb") as f:
        f.write(data)


async def _extract_text(file_bytes: bytes, filename: str) -> str:
    """按文件格式提取文本（在线程池执行以避免阻塞事件循环）。"""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

    if ext in ("txt", "md", "json", "csv"):
        # 文本格式直接解码
        return file_bytes.decode("utf-8", errors="replace")

    if ext == "pdf":
        return await asyncio.to_thread(_extract_pdf, file_bytes)

    if ext == "docx":
        return await asyncio.to_thread(_extract_docx, file_bytes)

    if ext == "xlsx":
        return await asyncio.to_thread(_extract_xlsx, file_bytes)

    # 未知格式 — 尝试当文本解码
    logger.warning("Unknown file format: %s, trying text decode", ext)
    return file_bytes.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    """pdfplumber 逐页提取文本。"""
    try:
        import pdfplumber

        with pdfplumber.open(io.BytesIO(data)) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
            return "\n".join(pages)
    except ImportError:
        logger.error("pdfplumber not installed, cannot parse PDF")
        return ""
    except Exception as e:
        logger.exception("PDF extraction failed")
        return ""


def _extract_docx(data: bytes) -> str:
    """python-docx 提取段落 + 表格文本。"""
    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                parts.append(" | ".join(cells))

        return "\n".join(parts)
    except ImportError:
        logger.error("python-docx not installed, cannot parse Word")
        return ""
    except Exception as e:
        logger.exception("DOCX extraction failed")
        return ""


def _extract_xlsx(data: bytes) -> str:
    """openpyxl 提取每个 sheet 的单元格文本。"""
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True)
        parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"--- Sheet: {sheet_name} ---")
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) for cell in row if cell is not None
                )
                if row_text.strip():
                    parts.append(row_text)

        wb.close()
        return "\n".join(parts)
    except ImportError:
        logger.error("openpyxl not installed, cannot parse Excel")
        return ""
    except Exception as e:
        logger.exception("XLSX extraction failed")
        return ""
